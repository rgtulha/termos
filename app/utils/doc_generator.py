import os
import datetime
import subprocess
from flask import current_app
from docx import Document
from copy import deepcopy 

class DocGenerator:
    """
    Classe responsável por manipular os templates .docx.
    (Corrigido para procurar placeholders em TODAS as tabelas)
    """
    def __init__(self, template_name):
        # ... (código __init__ não muda) ...
        self.template_name = template_name
        self.template_folder = current_app.config['TEMPLATE_FOLDER']
        self.output_folder = current_app.config['GENERATED_DOCS_FOLDER']
        self.template_path = os.path.join(self.template_folder, self.template_name)
        
        if not os.path.exists(self.template_path):
            current_app.logger.error(f"Arquivo de template não encontrado: {self.template_path}")
            raise FileNotFoundError(f"O template '{self.template_name}' não foi encontrado.")

    def generate_document(self, context, itens):
        """
        Gera o documento final.
        'context' = Dicionário com dados simples (ex: NOME_COMPLETO).
        'itens'   = Lista de objetos ItemEquipamento.
        """
        try:
            doc = Document(self.template_path)
            
            # 1. Preenche os placeholders simples (Nome, CPF, Data, etc.)
            self._replace_all_placeholders(doc, context)

            # 2. Preenche a tabela de equipamentos
            if itens:
                self._populate_table(doc, itens)

            # (O resto do processo de salvar e converter não muda)
            timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")
            base_filename = os.path.splitext(self.template_name)[0]
            docx_filename = f"{base_filename}_{timestamp}.docx"
            pdf_filename = f"{base_filename}_{timestamp}.pdf"
            docx_output_path = os.path.join(self.output_folder, docx_filename)
            
            doc.save(docx_output_path)
            
            current_app.logger.info(f"Iniciando conversão para PDF: {docx_output_path}")
            process = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                docx_output_path, '--outdir', self.output_folder
            ], timeout=15, capture_output=True, text=True)

            if process.returncode != 0:
                current_app.logger.error(f"Falha do LibreOffice: {process.stderr}")
                raise Exception(f"Falha ao converter para PDF: {process.stderr}")

            os.remove(docx_output_path)
            return pdf_filename

        except Exception as e:
            current_app.logger.error(f"Falha ao gerar o documento: {e}")
            raise Exception(f"Erro ao processar documento: {e}")


    # --- INÍCIO DA CORREÇÃO ---
    def _replace_all_placeholders(self, document, context):
        """
        Preenche placeholders SIMPLES (procura em parágrafos E tabelas).
        """
        # 1. Substituir no corpo principal (parágrafos)
        for p in document.paragraphs:
            self._replace_in_paragraph_robust(p, context)

        # 2. (ESTA ERA A PARTE EM FALTA)
        #    Substituir em tabelas no CORPO PRINCIPAL
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        # Verifica se o parágrafo não pertence à tabela de itens
                        # (para evitar substituições duplas)
                        if '{{item_seq}}' not in p.text and '{{item_descricao}}' not in p.text:
                             self._replace_in_paragraph_robust(p, context)
        
        # 3. Substituir em Cabeçalhos e Rodapés (Sem alteração)
        for section in document.sections:
            for p in section.header.paragraphs:
                self._replace_in_paragraph_robust(p, context)
            for t in section.header.tables:
                 for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_in_paragraph_robust(p, context)
            
            for p in section.footer.paragraphs:
                self._replace_in_paragraph_robust(p, context)
            for t in section.footer.tables:
                 for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_in_paragraph_robust(p, context)
    # --- FIM DA CORREÇÃO ---

    
    def _replace_in_paragraph_robust(self, paragraph, context):
        """
        Função de substituição robusta (preserva o estilo básico).
        """
        # ... (código não muda) ...
        full_text = paragraph.text
        if not any(f"{{{{{key}}}}}" in full_text for key, value in context.items()):
            return

        style = paragraph.runs[0].style if paragraph.runs else None
        
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"
            full_text = full_text.replace(placeholder, str(value or '')) 

        paragraph.clear()
        new_run = paragraph.add_run(full_text)
        if style:
            new_run.style = style
            
            
    def _populate_table(self, doc, itens):
        """
        Encontra a tabela de itens no documento e preenche-a.
        (Esta função está correta e não muda)
        """
        # ... (código não muda) ...
        target_table, template_row_index = None, -1

        for table in doc.tables:
            for i, row in enumerate(table.rows):
                cell_text = "".join([cell.text for cell in row.cells])
                if '{{item_seq}}' in cell_text or '{{item_descricao}}' in cell_text:
                    target_table = table
                    template_row_index = i
                    break
            if target_table:
                break
        
        if not target_table:
            current_app.logger.warning("Não foi encontrada nenhuma tabela de itens (com {{item_seq}} ou {{item_descricao}}) no template.")
            return

        template_row = target_table.rows[template_row_index]
        original_tr_xml = deepcopy(template_row._tr)

        for i, item in enumerate(itens):
            item_data_dict = {
                'item_seq': str(i + 1),
                'item_descricao': item.descricao,
                'item_patrimonio': item.patrimonio or '',
                'item_estado': item.estado or '',
            }
            if i == 0:
                row_to_fill = template_row
            else:
                new_tr = deepcopy(original_tr_xml)
                target_table._tbl.append(new_tr)
                row_to_fill = target_table.rows[-1]
            
            self._fill_row(row_to_fill, item_data_dict)

    def _fill_row(self, row, item_data_dict):
        """
        Preenche CADA célula de uma linha.
        (Esta função está correta e não muda)
        """
        # ... (código não muda) ...
        for cell in row.cells:
            for p in cell.paragraphs:
                if '{{' in p.text:
                    self._replace_in_paragraph_robust(p, item_data_dict)
