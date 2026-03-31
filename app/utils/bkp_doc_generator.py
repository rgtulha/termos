import os
import datetime
import subprocess
from flask import current_app
from docx import Document

class DocGenerator:
    """
    Classe responsável por manipular os templates .docx.
    (Corrigido o erro 'font.clone')
    """
    def __init__(self, template_name):
        self.template_name = template_name
        self.template_folder = current_app.config['TEMPLATE_FOLDER']
        self.output_folder = current_app.config['GENERATED_DOCS_FOLDER']
        self.template_path = os.path.join(self.template_folder, self.template_name)
        
        if not os.path.exists(self.template_path):
            current_app.logger.error(f"Arquivo de template não encontrado: {self.template_path}")
            raise FileNotFoundError(f"O template '{self.template_name}' não foi encontrado.")

    def generate_document(self, context):
        try:
            doc = Document(self.template_path)
            self._replace_all_placeholders(doc, context)

            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S%f")
            base_filename = os.path.splitext(self.template_name)[0]
            
            docx_filename = f"{base_filename}_{timestamp}.docx"
            pdf_filename = f"{base_filename}_{timestamp}.pdf"
            
            docx_output_path = os.path.join(self.output_folder, docx_filename)
            pdf_output_path = os.path.join(self.output_folder, pdf_filename)
            
            doc.save(docx_output_path)
            
            current_app.logger.info(f"Iniciando conversão para PDF: {docx_output_path}")

            process = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                docx_output_path, '--outdir', self.output_folder
            ], timeout=15, capture_output=True, text=True)

            if process.returncode != 0:
                current_app.logger.error(f"Falha do LibreOffice: {process.stderr}")
                raise Exception(f"Falha ao converter para PDF: {process.stderr}")

            current_app.logger.info(f"PDF gerado com sucesso: {pdf_output_path}")
            
            os.remove(docx_output_path)
            return pdf_filename

        except Exception as e:
            current_app.logger.error(f"Falha ao gerar o documento: {e}")
            raise Exception(f"Erro ao processar documento: {e}")


    def _replace_all_placeholders(self, document, context):
        """
        Função auxiliar ROBUSTA para varrer o documento (corpo, tabelas, cabeçalhos)
        e aplicar as substituições.
        """
        # 1. Substituir em tabelas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph_robust(p, context)
        
        # 2. Substituir no corpo principal (parágrafos)
        for p in document.paragraphs:
            self._replace_in_paragraph_robust(p, context)
        
        # 3. Substituir em Cabeçalhos e Rodapés
        for section in document.sections:
            for p in section.header.paragraphs:
                self._replace_in_paragraph_robust(p, context)
            for t in section.header.tables:
                 for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_in_paragraph_robust(p, context)
            
            for p in section.footer.paragraphs:
                self._replace_in_paragraph_robust(p, context)
            for t in section.footer.tables:
                 for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._replace_in_paragraph_robust(p, context)

    
    def _replace_in_paragraph_robust(self, paragraph, context):
        """
        Esta função substitui texto em um parágrafo inteiro,
        resolvendo o problema de 'runs' divididas.
        """
        full_text = paragraph.text
        if not any(f"{{{{{key}}}}}" in full_text for key, value in context.items()):
            return # Pula parágrafos que não têm placeholders

        # --- INÍCIO DA CORREÇÃO ---
        # Guarda o estilo do parágrafo/run. (Removida a linha 'font.clone()')
        style = paragraph.runs[0].style if paragraph.runs else None
        # --- FIM DA CORREÇÃO ---

        # Faz todas as substituições no texto completo
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"
            # Garante que o valor é uma string (e usa '' se for Nulo/None)
            full_text = full_text.replace(placeholder, str(value or '')) 

        # Limpa o parágrafo (remove todas as 'runs' antigas)
        paragraph.clear()

        # Adiciona o novo texto de volta em uma única 'run'
        new_run = paragraph.add_run(full_text)
        
        # --- INÍCIO DA CORREÇÃO ---
        # Reaplica apenas o estilo, que é seguro
        if style:
            new_run.style = style
        # --- FIM DA CORREÇÃO ---
